"""
docx_builder.py

Build a .docx file from structured content JSON.

Expected input format (same schema the extension produces):
{
  "blocks": [
    {
      "type": "paragraph",
      "runs": [
        {"type": "text",    "content": "The result is "},
        {"type": "formula", "latex": "x^2+y^2=r^2", "display": "inline"},
        {"type": "text",    "content": "."}
      ]
    },
    {
      "type": "formula_block",
      "latex": "\\\\int_0^\\\\infty e^{-x}\\,dx = 1"
    },
    {"type": "heading",   "level": 2, "text": "Section"},
    {"type": "code",      "text": "print('hello')"},
    {"type": "list_item", "runs": [...]}
  ]
}
"""

import io
import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from lxml import etree

from latex_to_omml import latex_to_omml_element

logger = logging.getLogger(__name__)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── low-level helpers ──────────────────────────────────────────────────────

def _add_text_run(para, text: str):
    """Append a plain text run to *para*."""
    if text:
        para.add_run(text)


def _add_formula(para, latex: str, display: str = "inline"):
    """
    Convert *latex* to OMML and append it to *para*._p.
    Falls back to italicised LaTeX text if conversion fails.
    """
    omml_el, err = latex_to_omml_element(latex, display)
    if err or omml_el is None:
        logger.warning("Formula fallback (plain text): %s  [%s]", latex, err)
        run = para.add_run(latex)
        run.italic = True
        return
    para._p.append(omml_el)


def _fill_runs(para, runs: list):
    """Fill *para* with a list of text/formula run dicts."""
    for run in runs:
        rtype = run.get("type", "text")
        if rtype == "text":
            _add_text_run(para, run.get("content", ""))
        elif rtype == "formula":
            _add_formula(para, run.get("latex", ""), run.get("display", "inline"))


# ── main builder ───────────────────────────────────────────────────────────

def build_docx(data: dict) -> bytes:
    """
    Build a .docx file from *data* and return its raw bytes.
    """
    doc = Document()

    # Remove the default empty paragraph that python-docx always adds
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)

    blocks = data.get("blocks", [])
    if not blocks:
        doc.add_paragraph("（内容为空）")

    for block in blocks:
        btype = block.get("type", "paragraph")

        # ── heading ───────────────────────────────────────────────────
        if btype == "heading":
            level = max(1, min(block.get("level", 1), 9))
            doc.add_heading(block.get("text", ""), level=level)

        # ── code block ────────────────────────────────────────────────
        elif btype == "code":
            para = doc.add_paragraph()
            run = para.add_run(block.get("text", ""))
            run.font.name = "Consolas"
            run.font.size = None  # inherit

        # ── standalone block formula ──────────────────────────────────
        elif btype == "formula_block":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_formula(para, block.get("latex", ""), "block")

        # ── table ─────────────────────────────────────────────────────
        elif btype == "table":
            rows_data = block.get("rows", [])
            if not rows_data:
                continue
            num_cols = max(len(r) for r in rows_data)
            table = doc.add_table(rows=0, cols=num_cols)
            table.style = "Table Grid"

            for row_cells in rows_data:
                row = table.add_row()
                for i, cell_data in enumerate(row_cells):
                    if i >= num_cols:
                        break
                    cell = row.cells[i]
                    # Clear default empty paragraph
                    cell.paragraphs[0].clear()
                    para = cell.paragraphs[0]
                    _fill_runs(para, cell_data.get("runs", []))

                    # Style header cells: bold + grey background
                    if cell_data.get("isHeader"):
                        for run in para.runs:
                            run.bold = True
                        shading = etree.SubElement(
                            cell._tc.get_or_add_tcPr(),
                            qn("w:shd"),
                        )
                        shading.set(qn("w:fill"), "F2F2F2")
                        shading.set(qn("w:val"), "clear")

        # ── list item ─────────────────────────────────────────────────
        elif btype == "list_item":
            try:
                para = doc.add_paragraph(style="List Bullet")
            except KeyError:
                para = doc.add_paragraph()
            _fill_runs(para, block.get("runs", []))

        # ── paragraph (default) ───────────────────────────────────────
        else:
            para = doc.add_paragraph()
            _fill_runs(para, block.get("runs", []))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
